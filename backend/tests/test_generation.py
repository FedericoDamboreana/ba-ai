from unittest.mock import patch


# Mock AI responses
MOCK_AI_QUESTIONS = {
    "questions": [
        {"question_text": "Test question 1?", "question_type": "Text", "is_critical": True},
        {"question_text": "Test question 2?", "question_type": "Text", "is_critical": True}
    ]
}

MOCK_AI_DOC = {
    "title": "Generated User Story",
    "user_story": {
        "as_a": "user",
        "i_want": "to login",
        "so_that": "I can access my account"
    },
    "acceptance_criteria": [
        {
            "scenario_name": "Successful login",
            "given": ["Given the user is on login page"],
            "when": ["When the user enters valid credentials"],
            "then": ["Then the user is logged in"]
        }
    ],
    "notes": "This is a test user story for login functionality",
    "dependencies": ["Authentication service", "User database"]
}

MOCK_KB = {"knowledge_base": "Updated knowledge base with new information"}


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_generate_documentation(mock_ai, client):
    """Test generating documentation with AI."""
    # Mock will be called 3 times: 1) questions, 2) doc generation, 3) knowledge base
    mock_ai.side_effect = [MOCK_AI_QUESTIONS, MOCK_AI_DOC, MOCK_KB]

    # Create project and item
    project_response = client.post(
        "/api/projects",
        json={"name": "Test Project", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "UserStory", "title": "Test Item", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Answer all critical questions
    questions_response = client.get(f"/api/items/{item_id}/questions")
    questions = questions_response.json()

    for question in questions:
        if question["is_critical"]:
            client.put(
                f"/api/questions/{question['id']}",
                json={"answer": "Test answer"}
            )

    # Generate documentation
    response = client.post(f"/api/items/{item_id}/generate", json={})
    assert response.status_code == 200
    data = response.json()
    assert data["item_id"] == item_id
    assert "content" in data
    assert data["content"]["title"] == "Generated User Story"
    assert "message" in data


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_regenerate_documentation(mock_ai, client):
    """Test regenerating documentation with feedback."""
    # Mock: 1) questions, 2) doc gen, 3) KB update, 4) regenerate
    mock_ai.side_effect = [MOCK_AI_QUESTIONS, MOCK_AI_DOC, MOCK_KB, MOCK_AI_DOC]

    # Create project and item
    project_response = client.post(
        "/api/projects",
        json={"name": "Test Project", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "PRD", "title": "Test PRD", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Answer questions
    questions_response = client.get(f"/api/items/{item_id}/questions")
    for question in questions_response.json():
        if question["is_critical"]:
            client.put(
                f"/api/questions/{question['id']}",
                json={"answer": "Test answer"}
            )

    # Initial generation
    client.post(f"/api/items/{item_id}/generate", json={})

    # Regenerate with feedback
    response = client.post(
        f"/api/items/{item_id}/regenerate",
        json={"feedback": "Make it more detailed"}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["item_id"] == item_id
    assert "content" in data


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_export_documentation(mock_ai, client):
    """Test exporting documentation as Word document."""
    from io import BytesIO
    from docx import Document

    # Mock: 1) questions, 2) doc gen, 3) KB update
    mock_ai.side_effect = [MOCK_AI_QUESTIONS, MOCK_AI_DOC, MOCK_KB]

    # Create project and item
    project_response = client.post(
        "/api/projects",
        json={"name": "Test Project", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "UserStory", "title": "Test Story", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Answer questions
    questions_response = client.get(f"/api/items/{item_id}/questions")
    for question in questions_response.json():
        if question["is_critical"]:
            client.put(
                f"/api/questions/{question['id']}",
                json={"answer": "Test answer"}
            )

    # Generate first to have content
    client.post(f"/api/items/{item_id}/generate", json={})

    # Export
    response = client.get(f"/api/items/{item_id}/export")
    assert response.status_code == 200

    # Check headers
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "content-disposition" in response.headers

    # Check filename format: project_name - documentation_name.docx
    content_disposition = response.headers["content-disposition"]
    assert "Test Project - Test Story.docx" in content_disposition

    # Check binary format
    assert b"PK" in response.content[:10]  # DOCX files start with PK (ZIP signature)

    # CRITICAL: Actually try to open the document with python-docx to verify it's valid
    try:
        buffer = BytesIO(response.content)
        doc = Document(buffer)

        # Verify the document has content
        assert len(doc.paragraphs) > 0, "Document has no paragraphs"

        # Check that key sections exist
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "User Story" in text_content, "Missing 'User Story' section"
        assert "Acceptance Criteria" in text_content, "Missing 'Acceptance Criteria' section"

        # Check the actual user story content is present
        assert "As a user" in text_content, "User story content missing"
        assert "I want to login" in text_content, "User story content missing"

        print("OK: Word document is valid and can be opened by python-docx")

    except Exception as e:
        raise AssertionError(f"Failed to open Word document with python-docx: {str(e)}")


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_export_with_special_characters_in_filename(mock_ai, client):
    """Test exporting documentation with special characters in project/item names."""
    from io import BytesIO
    from docx import Document
    from urllib.parse import unquote

    # Mock: 1) questions, 2) doc gen, 3) KB update
    mock_ai.side_effect = [MOCK_AI_QUESTIONS, MOCK_AI_DOC, MOCK_KB]

    # Create project and item with Spanish characters and spaces
    project_response = client.post(
        "/api/projects",
        json={"name": "Pedidos Ya - Partner Care", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "UserStory", "title": "Tablero de Tickets", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Answer questions
    questions_response = client.get(f"/api/items/{item_id}/questions")
    for question in questions_response.json():
        if question["is_critical"]:
            client.put(
                f"/api/questions/{question['id']}",
                json={"answer": "Test answer"}
            )

    # Generate first to have content
    client.post(f"/api/items/{item_id}/generate", json={})

    # Export
    response = client.get(f"/api/items/{item_id}/export")
    assert response.status_code == 200

    # Check headers
    content_disposition = response.headers["content-disposition"]

    # Should contain RFC 5987 filename* for UTF-8 encoding
    assert "filename*=UTF-8''" in content_disposition

    # Extract and decode the UTF-8 filename
    import re
    utf8_match = re.search(r"filename\*=UTF-8''([^;]+)", content_disposition)
    assert utf8_match is not None, "UTF-8 encoded filename not found"

    decoded_filename = unquote(utf8_match.group(1))
    assert "Pedidos Ya - Partner Care - Tablero de Tickets.docx" == decoded_filename

    # Verify the document is valid
    buffer = BytesIO(response.content)
    doc = Document(buffer)
    assert len(doc.paragraphs) > 0, "Document has no paragraphs"


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_export_with_invalid_filename_characters(mock_ai, client):
    """Test exporting documentation with characters invalid in filenames."""
    from io import BytesIO
    from docx import Document

    # Mock: 1) questions, 2) doc gen, 3) KB update
    mock_ai.side_effect = [MOCK_AI_QUESTIONS, MOCK_AI_DOC, MOCK_KB]

    # Create project and item with characters that are invalid in filenames
    project_response = client.post(
        "/api/projects",
        json={"name": "Project: Test/Name", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "UserStory", "title": "Item <Test> Name?", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Answer questions
    questions_response = client.get(f"/api/items/{item_id}/questions")
    for question in questions_response.json():
        if question["is_critical"]:
            client.put(
                f"/api/questions/{question['id']}",
                json={"answer": "Test answer"}
            )

    # Generate first to have content
    client.post(f"/api/items/{item_id}/generate", json={})

    # Export should succeed
    response = client.get(f"/api/items/{item_id}/export")
    assert response.status_code == 200

    # Verify invalid characters are sanitized (replaced with underscores)
    content_disposition = response.headers["content-disposition"]
    # Should not contain any of: < > : " / \ | ? *
    assert '<' not in content_disposition
    assert '>' not in content_disposition
    assert '?' not in content_disposition

    # Verify the document is still valid
    buffer = BytesIO(response.content)
    doc = Document(buffer)
    assert len(doc.paragraphs) > 0


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_export_prd_document(mock_ai, client):
    """Test exporting a PRD documentation type."""
    from io import BytesIO
    from docx import Document

    # PRD mock response
    mock_prd_doc = {
        "title": "Test PRD",
        "overview": "This is the project overview",
        "objectives": ["Objective 1", "Objective 2"],
        "scope": {
            "in_scope": ["Feature A", "Feature B"],
            "out_of_scope": ["Feature C"]
        },
        "stakeholders": [{"role": "Product Owner", "responsibilities": "Owns product vision"}],
        "requirements": [{"id": "REQ-001", "description": "Requirement 1", "priority": "High"}],
        "constraints": ["Budget constraint"],
        "success_criteria": ["Metric 1", "Metric 2"]
    }

    # Mock: 1) questions, 2) doc gen, 3) KB update
    mock_ai.side_effect = [MOCK_AI_QUESTIONS, mock_prd_doc, MOCK_KB]

    # Create project and PRD item
    project_response = client.post(
        "/api/projects",
        json={"name": "PRD Test Project", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "PRD", "title": "Test PRD Document", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Answer questions
    questions_response = client.get(f"/api/items/{item_id}/questions")
    for question in questions_response.json():
        if question["is_critical"]:
            client.put(
                f"/api/questions/{question['id']}",
                json={"answer": "Test answer"}
            )

    # Generate
    client.post(f"/api/items/{item_id}/generate", json={})

    # Export
    response = client.get(f"/api/items/{item_id}/export")
    assert response.status_code == 200

    # Verify the PRD document is valid and has expected content
    buffer = BytesIO(response.content)
    doc = Document(buffer)

    text_content = "\n".join([p.text for p in doc.paragraphs])
    assert "Overview" in text_content
    assert "Objectives" in text_content
    assert "Scope" in text_content
    assert "Requirements" in text_content


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_export_without_content(mock_ai, client):
    """Test exporting documentation without generated content."""
    mock_ai.return_value = MOCK_AI_QUESTIONS

    # Create project and item
    project_response = client.post(
        "/api/projects",
        json={"name": "Test Project", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "FRS", "title": "Test FRS", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Try to export without generating first
    response = client.get(f"/api/items/{item_id}/export")
    assert response.status_code == 400
    assert "No generated content" in response.json()["detail"]


def test_generate_nonexistent_item(client):
    """Test generating documentation for non-existent item."""
    response = client.post("/api/items/99999/generate", json={})
    assert response.status_code == 404


@patch('app.services.ai_service.ai_service.generate_structured_response')
def test_generate_without_questions_answered(mock_ai, client):
    """Test that generation fails when questions aren't answered."""
    mock_ai.return_value = MOCK_AI_QUESTIONS

    # Create project and item
    project_response = client.post(
        "/api/projects",
        json={"name": "Test Project", "description": "Test desc"}
    )
    project_id = project_response.json()["id"]

    item_response = client.post(
        f"/api/projects/{project_id}/items",
        json={"type": "UserStory", "title": "Test Item", "description": "Test desc"}
    )
    item_id = item_response.json()["id"]

    # Try to generate without answering questions
    response = client.post(f"/api/items/{item_id}/generate", json={})
    assert response.status_code == 400
    assert "Not all critical questions answered" in response.json()["detail"]
