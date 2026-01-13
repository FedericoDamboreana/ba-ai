"""
Tests for the export service - document generation functionality.
"""
from io import BytesIO
import pytest
from docx import Document
from app.services.export_service import (
    create_user_story_document,
    create_prd_document,
    export_to_word,
)


class TestUserStoryDocument:
    """Tests for User Story document generation."""

    def test_creates_valid_docx(self):
        """Test that create_user_story_document creates a valid docx file."""
        item_data = {
            "title": "Test User Story",
            "description": "Test description",
            "type": "UserStory"
        }
        generated_content = {
            "title": "Login Feature",
            "user_story": {
                "as_a": "registered user",
                "i_want": "to login to my account",
                "so_that": "I can access my dashboard"
            },
            "acceptance_criteria": [
                {
                    "scenario_name": "Successful Login",
                    "given": ["Given I am on the login page"],
                    "when": ["When I enter valid credentials"],
                    "then": ["Then I should be redirected to dashboard"]
                }
            ],
            "notes": "This is a critical feature",
            "dependencies": ["Authentication service"]
        }

        buffer = create_user_story_document(item_data, generated_content)

        # Verify buffer is not empty
        assert buffer.getvalue(), "Buffer is empty"

        # Verify it starts with ZIP signature (PK)
        buffer.seek(0)
        first_bytes = buffer.read(4)
        assert first_bytes[:2] == b'PK', "Not a valid ZIP/DOCX file"

        # Verify it can be opened by python-docx
        buffer.seek(0)
        doc = Document(buffer)

        # Verify content
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Login Feature" in text_content
        assert "As a registered user" in text_content
        assert "I want to login to my account" in text_content
        assert "Successful Login" in text_content

    def test_handles_empty_acceptance_criteria(self):
        """Test that document is created even with empty acceptance criteria."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Test Story",
            "user_story": {
                "as_a": "user",
                "i_want": "something",
                "so_that": "benefit"
            },
            "acceptance_criteria": []
        }

        buffer = create_user_story_document(item_data, generated_content)

        buffer.seek(0)
        doc = Document(buffer)
        assert len(doc.paragraphs) > 0

    def test_handles_missing_optional_fields(self):
        """Test that document is created when optional fields are missing."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Test Story",
            "user_story": {
                "as_a": "user",
                "i_want": "something",
                "so_that": "benefit"
            },
            "acceptance_criteria": []
            # notes and dependencies are missing
        }

        buffer = create_user_story_document(item_data, generated_content)

        buffer.seek(0)
        doc = Document(buffer)
        assert len(doc.paragraphs) > 0

    def test_handles_unicode_content(self):
        """Test that document handles Unicode characters correctly."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Gestión de Pedidos",
            "user_story": {
                "as_a": "usuario español",
                "i_want": "ver información en español",
                "so_that": "pueda entender mejor"
            },
            "acceptance_criteria": [
                {
                    "scenario_name": "Visualización correcta",
                    "given": ["Dado que estoy en la página"],
                    "when": ["Cuando selecciono español"],
                    "then": ["Entonces veo el contenido en español"]
                }
            ],
            "notes": "Caracteres especiales: áéíóú ñ ü",
            "dependencies": []
        }

        buffer = create_user_story_document(item_data, generated_content)

        buffer.seek(0)
        doc = Document(buffer)
        text_content = "\n".join([p.text for p in doc.paragraphs])

        assert "Gestión de Pedidos" in text_content
        assert "usuario español" in text_content
        assert "áéíóú ñ ü" in text_content


class TestPRDDocument:
    """Tests for PRD document generation."""

    def test_creates_valid_docx(self):
        """Test that create_prd_document creates a valid docx file."""
        item_data = {
            "title": "Test PRD",
            "description": "Test description",
            "type": "PRD"
        }
        generated_content = {
            "title": "Project Requirements",
            "overview": "This is the project overview",
            "objectives": ["Objective 1", "Objective 2"],
            "scope": {
                "in_scope": ["Feature A"],
                "out_of_scope": ["Feature B"]
            },
            "stakeholders": [{"role": "PM", "responsibilities": "Manage project"}],
            "requirements": [{"id": "REQ-001", "description": "Requirement", "priority": "High"}],
            "constraints": ["Budget limit"],
            "success_criteria": ["KPI 1"]
        }

        buffer = create_prd_document(item_data, generated_content)

        buffer.seek(0)
        doc = Document(buffer)

        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Project Requirements" in text_content
        assert "Overview" in text_content
        assert "Objectives" in text_content
        assert "In Scope" in text_content
        assert "Out of Scope" in text_content

    def test_handles_empty_lists(self):
        """Test that PRD document is created with empty lists."""
        item_data = {"title": "Test", "description": "Test", "type": "PRD"}
        generated_content = {
            "title": "Empty PRD",
            "overview": "Overview",
            "objectives": [],
            "scope": {"in_scope": [], "out_of_scope": []},
            "stakeholders": [],
            "requirements": [],
            "constraints": [],
            "success_criteria": []
        }

        buffer = create_prd_document(item_data, generated_content)

        buffer.seek(0)
        doc = Document(buffer)
        assert len(doc.paragraphs) > 0


class TestExportToWord:
    """Tests for the main export_to_word function."""

    def test_routes_to_user_story(self):
        """Test that UserStory type routes to user story document."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Story",
            "user_story": {"as_a": "user", "i_want": "x", "so_that": "y"},
            "acceptance_criteria": []
        }

        buffer = export_to_word(item_data, generated_content, "UserStory")

        buffer.seek(0)
        doc = Document(buffer)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "User Story" in text_content

    def test_routes_to_prd(self):
        """Test that PRD type routes to PRD document."""
        item_data = {"title": "Test", "description": "Test", "type": "PRD"}
        generated_content = {
            "title": "PRD",
            "overview": "Overview",
            "objectives": [],
            "scope": {"in_scope": [], "out_of_scope": []},
            "requirements": []
        }

        buffer = export_to_word(item_data, generated_content, "PRD")

        buffer.seek(0)
        doc = Document(buffer)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Overview" in text_content

    def test_handles_epic_type(self):
        """Test that Epic type creates a generic document."""
        item_data = {"title": "Test Epic", "description": "Test", "type": "Epic"}
        generated_content = {"title": "Epic Title", "features": ["Feature 1"]}

        buffer = export_to_word(item_data, generated_content, "Epic")

        buffer.seek(0)
        doc = Document(buffer)
        assert len(doc.paragraphs) > 0

    def test_handles_frs_type(self):
        """Test that FRS type creates a generic document."""
        item_data = {"title": "Test FRS", "description": "Test", "type": "FRS"}
        generated_content = {"title": "FRS Title", "requirements": []}

        buffer = export_to_word(item_data, generated_content, "FRS")

        buffer.seek(0)
        doc = Document(buffer)
        assert len(doc.paragraphs) > 0

    def test_buffer_is_seeked_to_start(self):
        """Test that the buffer is properly seeked to the start."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Story",
            "user_story": {"as_a": "user", "i_want": "x", "so_that": "y"},
            "acceptance_criteria": []
        }

        buffer = export_to_word(item_data, generated_content, "UserStory")

        # Buffer should be at position 0
        assert buffer.tell() == 0

        # First bytes should be PK (ZIP signature)
        first_bytes = buffer.read(2)
        assert first_bytes == b'PK'


class TestDocumentIntegrity:
    """Tests to ensure documents are not corrupted."""

    def test_document_can_be_read_multiple_times(self):
        """Test that the document can be read from buffer multiple times."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Story",
            "user_story": {"as_a": "user", "i_want": "x", "so_that": "y"},
            "acceptance_criteria": []
        }

        buffer = export_to_word(item_data, generated_content, "UserStory")

        # Read the content
        content = buffer.getvalue()

        # Create new buffers from the same content and verify they work
        for _ in range(3):
            new_buffer = BytesIO(content)
            doc = Document(new_buffer)
            assert len(doc.paragraphs) > 0

    def test_large_document(self):
        """Test that large documents are created correctly."""
        item_data = {"title": "Test", "description": "Test", "type": "UserStory"}
        generated_content = {
            "title": "Large User Story",
            "user_story": {
                "as_a": "user",
                "i_want": "many features " * 100,
                "so_that": "many benefits " * 100
            },
            "acceptance_criteria": [
                {
                    "scenario_name": f"Scenario {i}",
                    "given": [f"Given condition {j}" for j in range(10)],
                    "when": [f"When action {j}" for j in range(5)],
                    "then": [f"Then result {j}" for j in range(10)]
                }
                for i in range(20)
            ],
            "notes": "Long notes " * 500,
            "dependencies": [f"Dependency {i}" for i in range(50)]
        }

        buffer = export_to_word(item_data, generated_content, "UserStory")

        # Verify the large document is valid
        buffer.seek(0)
        doc = Document(buffer)
        assert len(doc.paragraphs) > 100
